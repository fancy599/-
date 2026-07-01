from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[3]
SOURCE_ROOT = (
    Path.home()
    / "Desktop"
    / "\u5929\u804c\u9879\u76ee"
    / "202606-\u4e34\u6e2f\u63a7\u80a1\u5236\u5ea6\u4fee\u8ba2"
    / "02\u3001\u4e0a\u6d77\u5e02\u56fd\u8d44\u59d4"
    / "02\u3001\u4e0a\u6d77\u5e02\u56fd\u8d44\u59d4"
)
OUT_DIR = ROOT / "ic-review-app" / "backend" / "app" / "data" / "regulatory_corpus"
OUT_PATH = OUT_DIR / "regulatory_corpus.json"

SUPPORTED = {".pdf", ".docx", ".xlsx"}
MAX_TEXT_CHARS = 260_000


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    import fitz

    doc = fitz.open(str(path))
    parts: list[str] = []
    try:
        for page in doc:
            text = page.get_text("text", sort=True).strip()
            if text:
                parts.append(f"[PAGE {page.number + 1}]\n{text}")
        meta = {"page_count": doc.page_count}
    finally:
        doc.close()
    return clean_text("\n\n".join(parts)), meta


def extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            values = [" ".join(cell.text.split()) for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            parts.append(f"[TABLE {table_index}]")
            parts.extend(rows)
    return clean_text("\n".join(parts)), {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}


def extract_xlsx(path: Path) -> tuple[str, dict[str, Any]]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            parts.append(f"[SHEET] {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = [str(v).strip() for v in row if v not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
    finally:
        wb.close()
    return clean_text("\n".join(parts)), {"sheet_count": len(wb.sheetnames)}


def paragraphs(text: str) -> list[str]:
    chunks = []
    for raw in re.split(r"\n{2,}|\n(?=(?:第[一二三四五六七八九十百零〇\d]+[章节条款]|[一二三四五六七八九十]+、|[（(][一二三四五六七八九十\d]+[）)]))", text):
        item = " ".join(raw.split())
        if len(item) >= 12:
            chunks.append(item[:1200])
    return chunks


def extract(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            text, meta = extract_pdf(path)
        elif suffix == ".docx":
            text, meta = extract_docx(path)
        elif suffix == ".xlsx":
            text, meta = extract_xlsx(path)
        else:
            return {"ok": False, "error": f"unsupported suffix {suffix}"}
        text = text[:MAX_TEXT_CHARS]
        paras = paragraphs(text)
        return {
            "ok": True,
            "relative_path": str(path.relative_to(SOURCE_ROOT)),
            "file_name": path.name,
            "suffix": suffix,
            "size": path.stat().st_size,
            "meta": meta,
            "text_len": len(text),
            "title_guess": next((p for p in paras[:8] if len(p) <= 120), path.stem),
            "paragraphs": paras,
        }
    except Exception as exc:
        return {
            "ok": False,
            "relative_path": str(path.relative_to(SOURCE_ROOT)),
            "file_name": path.name,
            "suffix": suffix,
            "size": path.stat().st_size,
            "error": f"{type(exc).__name__}: {exc}",
        }


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        p
        for p in SOURCE_ROOT.rglob("*")
        if p.is_file()
        and p.suffix.lower() in SUPPORTED
        and not p.name.startswith("~$")
        and p.name != "\u30100204\u3011\u5916\u89c4\u6c47\u7f16.zip"
    ]
    records = [extract(path) for path in sorted(files, key=lambda p: str(p.relative_to(SOURCE_ROOT)))]
    OUT_PATH.write_text(json.dumps({"source_root": str(SOURCE_ROOT), "records": records}, ensure_ascii=False, indent=2), encoding="utf-8")
    ok = sum(1 for r in records if r.get("ok"))
    print(json.dumps({"files": len(records), "ok": ok, "failed": len(records) - ok, "out": str(OUT_PATH)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
