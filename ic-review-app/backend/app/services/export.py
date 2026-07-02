from collections import Counter
from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy.orm import Session

from app.models import Difference, ReviewLog, ReviewTask
from app.utils.timeutil import to_local_iso

# 风险等级展示顺序，保证“高风险优先”聚合（与 PRD 8.5/8.8 一致）。
_RISK_ORDER = {"高": 0, "中": 1, "低": 2}
_ACTION_LABEL = {"confirmed": "已确认", "rejected": "已驳回", "need_follow_up": "需跟进"}


def _review_status_label(status: str) -> str:
    return {
        "pending": "待复核",
        "pending_evidence": "待补充证据",
        "confirmed": "已确认",
        "rejected": "已驳回",
        "need_follow_up": "需跟进",
    }.get(status, status or "")


def _sorted_diffs(db: Session, task: ReviewTask) -> list[Difference]:
    diffs = db.query(Difference).filter(Difference.task_id == task.id).all()
    return sorted(
        diffs,
        key=lambda d: (_RISK_ORDER.get(d.risk_level, 9), d.id),
    )


def _report_stats(diffs: list[Difference]) -> dict:
    risk_counter = Counter(d.risk_level for d in diffs)
    type_counter = Counter(d.diff_type for d in diffs)
    return {
        "total": len(diffs),
        "high": risk_counter.get("高", 0),
        "mid": risk_counter.get("中", 0),
        "low": risk_counter.get("低", 0),
        "pending": sum(1 for d in diffs if d.review_status in ("pending", "pending_evidence")),
        "confirmed": sum(1 for d in diffs if d.review_status == "confirmed"),
        "evidence_pending": sum(1 for d in diffs if not d.evidence_ok),
        "risk_counter": risk_counter,
        "type_counter": type_counter,
    }


# --------------------------------------------------------------------------- #
# Excel
# --------------------------------------------------------------------------- #
def export_task_xlsx(db: Session, task: ReviewTask) -> bytes:
    diffs = _sorted_diffs(db, task)
    logs = db.query(ReviewLog).filter(ReviewLog.task_id == task.id).order_by(ReviewLog.id.asc()).all()
    stats = _report_stats(diffs)

    wb = Workbook()
    header_fill = PatternFill("solid", fgColor="22262B")
    header_font = Font(color="FFFFFF", bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")

    def _style_header(ws) -> None:
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(vertical="center")

    # Sheet 1 概览与统计
    ws_overview = wb.active
    ws_overview.title = "概览"
    ws_overview.append(["项目", "内容"])
    overview_rows = [
        ("任务名称", task.task_name),
        ("业务领域", task.business_domain),
        ("任务类型", "单制度体检" if getattr(task, "task_type", "") == "single" else "集团对子公司"),
        ("任务状态", task.status),
        ("差异总数", stats["total"]),
        ("高风险", stats["high"]),
        ("中风险", stats["mid"]),
        ("低风险", stats["low"]),
        ("待复核", stats["pending"]),
        ("已确认", stats["confirmed"]),
        ("待补充证据", stats["evidence_pending"]),
        ("差异类型分布", "；".join(f"{k}:{v}" for k, v in stats["type_counter"].most_common())),
        ("生成时间", to_local_iso(task.completed_at or task.created_at)),
        ("审查说明", task.report_summary or ""),
    ]
    for row in overview_rows:
        ws_overview.append(list(row))
    _style_header(ws_overview)
    ws_overview.column_dimensions["A"].width = 16
    ws_overview.column_dimensions["B"].width = 80
    for row in ws_overview.iter_rows(min_row=2):
        row[1].alignment = wrap

    # Sheet 2 差异明细（含原文证据、AI 理由、复核结论）
    ws = wb.create_sheet("差异明细")
    headers = [
        "风险", "差异类型", "控制主题", "问题摘要",
        "集团/标准依据原文", "子公司/本制度原文",
        "集团位置", "子公司位置",
        "AI 判断理由", "修改建议", "置信度", "证据校验", "复核状态",
    ]
    ws.append(headers)
    for d in diffs:
        ws.append([
            d.risk_level, d.diff_type, d.control_topic, d.summary,
            d.group_excerpt, d.subsidiary_excerpt,
            d.group_location, d.subsidiary_location,
            d.ai_reason, d.suggestion,
            round(d.confidence or 0.0, 2),
            "通过" if d.evidence_ok else "待补证",
            _review_status_label(d.review_status),
        ])
    _style_header(ws)
    widths = [6, 12, 16, 34, 40, 40, 18, 18, 44, 40, 8, 10, 12]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = wrap

    # Sheet 3 高风险汇总
    ws_high = wb.create_sheet("高风险汇总")
    ws_high.append(["控制主题", "差异类型", "问题摘要", "修改建议", "复核状态"])
    for d in diffs:
        if d.risk_level == "高":
            ws_high.append([d.control_topic, d.diff_type, d.summary, d.suggestion, _review_status_label(d.review_status)])
    _style_header(ws_high)
    for i, w in enumerate([18, 12, 40, 40, 12], start=1):
        ws_high.column_dimensions[ws_high.cell(row=1, column=i).column_letter].width = w
    for row in ws_high.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = wrap

    # Sheet 4 复核记录（复核链）
    ws_log = wb.create_sheet("复核记录")
    ws_log.append(["复核人", "操作", "复核意见", "时间"])
    for log in logs:
        ws_log.append([
            log.reviewer,
            _ACTION_LABEL.get(log.action, log.action),
            log.comment,
            to_local_iso(log.created_at),
        ])
    _style_header(ws_log)
    for i, w in enumerate([16, 10, 50, 20], start=1):
        ws_log.column_dimensions[ws_log.cell(row=1, column=i).column_letter].width = w
    for row in ws_log.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = wrap

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Word（PRD 8.8 要求的 Word 底稿）
# --------------------------------------------------------------------------- #
def export_task_docx(db: Session, task: ReviewTask) -> bytes:
    diffs = _sorted_diffs(db, task)
    logs = db.query(ReviewLog).filter(ReviewLog.task_id == task.id).order_by(ReviewLog.id.asc()).all()
    stats = _report_stats(diffs)
    is_single = getattr(task, "task_type", "") == "single"
    left_label = "标准依据 / 集团原文" if not is_single else "标准控制点依据"
    right_label = "子公司原文覆盖" if not is_single else "本制度覆盖情况"

    doc = DocxDocument()
    title = doc.add_heading("内控制度智能审查底稿", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 任务信息
    doc.add_heading("一、任务信息", level=1)
    info = doc.add_table(rows=0, cols=2)
    info.style = "Light Grid Accent 1"
    for k, v in [
        ("任务名称", task.task_name),
        ("业务领域", task.business_domain),
        ("任务类型", "单制度体检" if is_single else "集团对子公司"),
        ("任务状态", task.status),
        ("生成时间", to_local_iso(task.completed_at or task.created_at)),
    ]:
        cells = info.add_row().cells
        cells[0].text = k
        cells[1].text = str(v or "")

    # 2. 差异统计
    doc.add_heading("二、差异统计", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"共识别差异 {stats['total']} 条，其中高风险 {stats['high']} 条、"
        f"中风险 {stats['mid']} 条、低风险 {stats['low']} 条；"
        f"待复核 {stats['pending']} 条，待补充证据 {stats['evidence_pending']} 条。"
    )
    if stats["type_counter"]:
        doc.add_paragraph("差异类型分布：" + "；".join(f"{k} {v} 条" for k, v in stats["type_counter"].most_common()))
    if task.report_summary:
        doc.add_paragraph("审查摘要：" + task.report_summary)

    # 3. 高风险汇总
    doc.add_heading("三、高风险汇总", level=1)
    high_diffs = [d for d in diffs if d.risk_level == "高"]
    if high_diffs:
        ht = doc.add_table(rows=1, cols=3)
        ht.style = "Light Grid Accent 1"
        hdr = ht.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "控制主题", "差异类型", "问题摘要"
        for d in high_diffs:
            c = ht.add_row().cells
            c[0].text = d.control_topic or ""
            c[1].text = d.diff_type or ""
            c[2].text = d.summary or ""
    else:
        doc.add_paragraph("无高风险差异。")

    # 4. 差异明细（含原文证据、AI 理由、修改建议、复核结论）
    doc.add_heading("四、差异明细", level=1)
    for idx, d in enumerate(diffs, start=1):
        h = doc.add_heading(level=2)
        run = h.add_run(f"{idx}. [{d.risk_level}风险/{d.diff_type}] {d.control_topic or d.summary}")
        if d.risk_level == "高":
            run.font.color.rgb = RGBColor(0xA8, 0x3F, 0x39)
        doc.add_paragraph(d.summary or "")

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light List Accent 1"
        for k, v in [
            (left_label, f"{d.group_excerpt or ''}\n（位置：{d.group_location or '—'}）"),
            (right_label, f"{d.subsidiary_excerpt or ''}\n（位置：{d.subsidiary_location or '—'}）"),
            ("AI 判断理由", d.ai_reason or ""),
            ("修改建议", d.suggestion or ""),
            ("置信度 / 证据校验", f"{round(d.confidence or 0.0, 2)} / {'通过' if d.evidence_ok else '待补证'}"),
            ("复核状态", _review_status_label(d.review_status)),
        ]:
            cells = tbl.add_row().cells
            cells[0].text = k
            cells[0].paragraphs[0].runs[0].font.bold = True if cells[0].paragraphs[0].runs else False
            cells[1].text = v

    # 5. 人工复核结论与复核链
    doc.add_heading("五、人工复核结论与复核链", level=1)
    if logs:
        lt = doc.add_table(rows=1, cols=4)
        lt.style = "Light Grid Accent 1"
        hdr = lt.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "复核人", "操作", "复核意见", "时间"
        for log in logs:
            c = lt.add_row().cells
            c[0].text = log.reviewer or ""
            c[1].text = _ACTION_LABEL.get(log.action, log.action)
            c[2].text = log.comment or ""
            c[3].text = to_local_iso(log.created_at)
    else:
        doc.add_paragraph("暂无人工复核记录。")

    foot = doc.add_paragraph("内控制度智能审查平台 · 自动生成底稿，结论需经三级复核签字后方可正式出具。")
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(0x70, 0x77, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# HTML
# --------------------------------------------------------------------------- #
def export_task_html(db: Session, task: ReviewTask) -> str:
    diffs = _sorted_diffs(db, task)
    logs = db.query(ReviewLog).filter(ReviewLog.task_id == task.id).order_by(ReviewLog.id.asc()).all()
    stats = _report_stats(diffs)

    def esc(v: object) -> str:
        return (
            str(v or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br>")
        )

    rows = "".join(
        f"<tr>"
        f"<td><span class='pill {('red' if d.risk_level == '高' else 'amber' if d.risk_level == '中' else 'blue')}'>{esc(d.risk_level)}</span></td>"
        f"<td>{esc(d.diff_type)}</td><td>{esc(d.control_topic)}</td><td>{esc(d.summary)}</td>"
        f"<td class='ev'>{esc(d.group_excerpt)}</td><td class='ev'>{esc(d.subsidiary_excerpt)}</td>"
        f"<td class='ev'>{esc(d.ai_reason)}</td><td class='ev'>{esc(d.suggestion)}</td>"
        f"<td>{'通过' if d.evidence_ok else '待补证'}</td>"
        f"<td>{esc(_review_status_label(d.review_status))}</td></tr>"
        for d in diffs
    )
    log_rows = "".join(
        f"<tr><td>{esc(log.reviewer)}</td><td>{esc(_ACTION_LABEL.get(log.action, log.action))}</td>"
        f"<td>{esc(log.comment)}</td><td class='muted'>{esc(to_local_iso(log.created_at))}</td></tr>"
        for log in logs
    )
    summary = esc(task.report_summary or "（无摘要）")
    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{esc(task.task_name)} · 审查报告</title>
<style>
:root {{
  --bg:#f7f7f5;--panel:#fff;--ink:#22262b;--muted:#707780;--line:#dfe3e8;--soft:#eceff2;
  --brand:#252b33;--amber:#d98513;--red:#a83f39;--ok:#2d6a4f;--done:#3f5f83;
  --font:"Microsoft YaHei","PingFang SC",sans-serif;
}}
*{{box-sizing:border-box}} body{{margin:0;font-family:var(--font);font-size:14px;color:var(--ink);
background:linear-gradient(180deg,#fbfbfa 0%,var(--bg) 200px);line-height:1.65}}
.wrap{{max-width:1180px;margin:0 auto;padding:32px 24px 48px}}
h1{{margin:0 0 8px;font-size:22px;font-weight:760}}
.sub{{color:var(--muted);font-size:13px;margin:0 0 24px}}
.stats{{display:flex;flex-wrap:wrap;gap:16px 28px;margin:20px 0;padding:16px 18px;
background:#fff;border:1px solid var(--soft);border-radius:8px;font-weight:700;font-size:13px}}
.panel{{background:var(--panel);border:1px solid var(--soft);border-radius:8px;
box-shadow:0 1px 2px rgba(28,34,43,.04);margin-bottom:20px;overflow:hidden}}
.panel-hd{{padding:14px 18px;border-bottom:1px solid var(--soft);font-weight:760;font-size:15px}}
table{{border-collapse:collapse;width:100%}}
th,td{{border-bottom:1px solid var(--soft);padding:11px 12px;text-align:left;vertical-align:top}}
th{{background:#fafafa;color:var(--muted);font-size:12px;font-weight:700}}
td.ev{{font-size:12.5px;color:#3f4a58;max-width:240px}}
.pill{{display:inline-block;padding:2px 8px;border-radius:5px;font-size:12px;font-weight:700}}
.pill.red{{background:#f8edeb;color:var(--red)}} .pill.amber{{background:#fff4db;color:#8d6118}}
.pill.blue{{background:#edf3f8;color:var(--done)}} .muted{{color:var(--muted);font-size:12px}}
.summary{{background:#edf7f1;border:1px solid #bddbd7;border-radius:8px;padding:16px 18px;margin-bottom:20px}}
.summary h2{{margin:0 0 8px;font-size:15px;color:var(--ok)}}
.footer{{margin-top:32px;font-size:12px;color:var(--muted);text-align:center}}
</style></head>
<body><div class="wrap">
<h1>{esc(task.task_name)}</h1>
<p class="sub">内控制度智能审查报告 · 生成时间 {esc(to_local_iso(task.completed_at or task.created_at))}</p>
<div class="stats">
  <span>差异：{stats['total']} 条</span>
  <span>高风险：{stats['high']} 条</span>
  <span>中风险：{stats['mid']} 条</span>
  <span>低风险：{stats['low']} 条</span>
  <span>待复核：{stats['pending']} 条</span>
</div>
<div class="summary"><h2>审查摘要</h2><p style="margin:0">{summary}</p></div>
<div class="panel"><div class="panel-hd">差异清单</div>
<table><thead><tr><th>风险</th><th>类型</th><th>主题</th><th>摘要</th><th>集团/标准依据</th><th>子公司/本制度原文</th><th>AI 理由</th><th>修改建议</th><th>证据</th><th>复核</th></tr></thead>
<tbody>{rows or "<tr><td colspan='10' class='muted'>暂无差异</td></tr>"}</tbody></table></div>
<div class="panel"><div class="panel-hd">复核记录</div>
<table><thead><tr><th>复核人</th><th>操作</th><th>意见</th><th>时间</th></tr></thead>
<tbody>{log_rows or "<tr><td colspan='4' class='muted'>暂无复核记录</td></tr>"}</tbody></table></div>
<p class="footer">内控制度智能审查平台 · 结论需经三级复核签字后方可正式出具</p>
</div></body></html>"""
