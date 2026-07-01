import shutil
import subprocess
import sys
import tempfile
import os
import re
import json
import logging
import time
from pathlib import Path

import fitz

from app.config import ROOT_DIR, get_settings
from app.services.text_normalize import cjk_ratio, clean_page_body, normalize_extracted_text

ALLOWED_SUFFIXES = {".txt", ".doc", ".docx", ".pdf"}
logger = logging.getLogger(__name__)


class ParseError(Exception):
    def __init__(self, message: str, code: str = "ERR_DOC_PARSE_FAILED"):
        self.code = code
        self.message = message
        super().__init__(f"{code}: {message}")


def _debug_log(run_id: str, hypothesis_id: str, location: str, message: str, data: dict) -> None:
    logger.debug("%s %s %s %s %s", run_id, hypothesis_id, location, message, data)


def extract_text_from_file(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise ParseError(f"文件不存在: {file_path}", "ERR_DOC_NOT_FOUND")

    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".doc":
        text = _extract_doc(path)
    elif suffix == ".pdf":
        text = _extract_pdf_with_provider(path)
    else:
        raise ParseError(f"不支持的文件格式: {suffix}，支持 .txt / .doc / .docx / .pdf", "ERR_DOC_UNSUPPORTED")

    text = text.strip()
    # region agent log
    _debug_log(
        "pre-fix",
        "H1",
        "parser.py:extract_text_from_file",
        "raw extraction result",
        {"suffix": suffix, "raw_len": len(text), "raw_cjk_ratio": round(cjk_ratio(text), 4)},
    )
    # endregion
    if not text:
        raise ParseError(
            "未能提取到文字内容，请确认文件非空，或上传脱密版、可编辑版文档",
            "ERR_DOC_EMPTY_OR_ENCRYPTED",
        )
    normalized = normalize_extracted_text(text)
    # region agent log
    _debug_log(
        "pre-fix",
        "H1",
        "parser.py:extract_text_from_file",
        "normalized extraction result",
        {"suffix": suffix, "normalized_len": len(normalized), "normalized_cjk_ratio": round(cjk_ratio(normalized), 4)},
    )
    # endregion
    return normalized


def detect_pdf_complex_table_pages(file_path: str) -> list[int]:
    """Best-effort quality gate. It flags likely complex table pages; it never claims exact parsing."""
    path = Path(file_path)
    if path.suffix.lower() != ".pdf":
        return []
    try:
        doc = fitz.open(str(path))
    except Exception:
        return []
    pages: list[int] = []
    try:
        for index, page in enumerate(doc):
            text = page.get_text("text") or ""
            drawings = page.get_drawings()
            line_like = sum(1 for drawing in drawings if len(drawing.get("items", [])) >= 2)
            table_terms = sum(text.count(term) for term in ("审批", "金额", "权限", "党委会", "董事会", "责任部门"))
            if (line_like >= 8 and table_terms >= 2) or line_like >= 18:
                pages.append(index + 1)
    finally:
        doc.close()
    return pages


def render_pdf_page_snapshot(file_path: str, page_number: int) -> bytes:
    path = Path(file_path)
    if path.suffix.lower() != ".pdf":
        raise ParseError("仅 PDF 支持页面快照", "ERR_SNAPSHOT_UNSUPPORTED")
    doc = fitz.open(str(path))
    try:
        if page_number < 1 or page_number > doc.page_count:
            raise ParseError("页码超出范围", "ERR_SNAPSHOT_PAGE")
        pix = doc[page_number - 1].get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        return pix.tobytes("png")
    finally:
        doc.close()


def _extract_doc(path: Path) -> str:
    """旧版 .doc：Windows 优先 Word，其次 LibreOffice 转 docx 再解析。"""
    errors: list[str] = []

    if sys.platform == "win32":
        try:
            return _extract_doc_via_word_com(path)
        except Exception as e:
            errors.append(f"Word 解析: {e}")

    try:
        return _extract_doc_via_libreoffice(path)
    except Exception as e:
        errors.append(f"LibreOffice 转换: {e}")

    hint = (
        "无法解析 .doc 文件。"
        " Windows 请安装 Microsoft Word，或安装 LibreOffice；"
        " 也可在 Word 中另存为 .docx 后上传。"
    )
    if errors:
        hint += " 详情：" + "；".join(errors)
    raise ParseError(hint)


def _extract_doc_via_word_com(path: Path) -> str:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    doc = None
    tmp_dir = tempfile.mkdtemp(prefix="ic_doc_")
    out_docx = Path(tmp_dir) / "converted.docx"
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        abs_path = str(path.resolve())
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        # 16 = wdFormatXMLDocument (.docx)
        doc.SaveAs2(str(out_docx.resolve()), FileFormat=16)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        if not out_docx.exists():
            raise ParseError("Word 转换 docx 失败")
        return _extract_docx(out_docx)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        shutil.rmtree(tmp_dir, ignore_errors=True)
        pythoncom.CoUninitialize()


def _extract_doc_via_libreoffice(path: Path) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise ParseError("未找到 LibreOffice（soffice）")

    tmp_dir = Path(tempfile.mkdtemp(prefix="ic_doc_"))
    try:
        proc = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_dir),
                str(path.resolve()),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or "").strip()
            raise ParseError(err or f"LibreOffice 退出码 {proc.returncode}")

        candidates = list(tmp_dir.glob("*.docx"))
        if not candidates:
            raise ParseError("LibreOffice 未生成 docx 文件")
        return _extract_docx(candidates[0])
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _extract_docx_raw_xml(path: Path) -> str:
    """兜底：直接从 OOXML 抽取所有文字（含文本框 txbxContent、页眉页脚）。

    python-docx 的 doc.paragraphs 读不到文本框/图形里的文字，很多国企/WPS 制度正文恰好
    放在文本框里，导致解析出近乎空白。本函数对 document.xml 及页眉页脚做去标签处理，
    把段落结束转换行，保留所有 <w:t> 文本。"""
    import html
    import re as _re
    import zipfile

    try:
        zf = zipfile.ZipFile(str(path))
    except Exception:
        return ""
    targets = [
        n
        for n in zf.namelist()
        if n == "word/document.xml" or _re.match(r"word/(header|footer)\d*\.xml$", n)
    ]
    chunks: list[str] = []
    for name in targets:
        try:
            xml = zf.read(name).decode("utf-8", "ignore")
        except Exception:
            continue
        xml = _re.sub(r"</w:p>", "\n", xml)
        xml = _re.sub(r"<w:(?:tab|br|cr)\b[^>]*/?>", " ", xml)
        plain = _re.sub(r"<[^>]+>", "", xml)
        chunks.append(html.unescape(plain))
    zf.close()
    out = "\n".join(chunks)
    out = _re.sub(r"[ \t]+\n", "\n", out)
    out = _re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _extract_docx(path: Path) -> str:
    parts: list[str] = []
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ParseError("python-docx 未安装") from e

    try:
        doc = DocxDocument(str(path))
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table_index, table in enumerate(doc.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                values = [normalize_extracted_text(cell.text).replace("\n", " / ").strip() for cell in row.cells]
                if any(values):
                    rows.append(values)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            padded = [row + [""] * (width - len(row)) for row in rows]
            parts.append(f"[TABLE {table_index}]")
            parts.append("| " + " | ".join(padded[0]) + " |")
            parts.append("| " + " | ".join(["---"] * width) + " |")
            for row in padded[1:]:
                parts.append("| " + " | ".join(row) + " |")
            parts.append(f"[/TABLE {table_index}]")
    except Exception as e:
        # python-docx 打不开（WPS/异常结构等）：留给下面的 OOXML 兜底，仍打不开才报错。
        logger.warning("python-docx 解析失败，尝试 OOXML 兜底：%s（%s）", path.name, e)

    pydocx_text = "\n".join(parts)
    raw_text = _extract_docx_raw_xml(path)
    # 取信息量更大的一份：文本框正文常常只有 OOXML 兜底能拿到。
    best = pydocx_text if len(pydocx_text) >= len(raw_text) else raw_text

    if not best.strip():
        raise ParseError(
            "Word 中未找到可识别文字（可能正文为图片/扫描、文件已加密，或为旧版 .doc 误改成 .docx）",
            "ERR_DOC_EMPTY_OR_ENCRYPTED",
        )
    return best


def _extract_pdf_with_provider(path: Path) -> str:
    """PDF 解析：配置百度·文档解析时优先用模型解析，失败或为空回退本地 PyMuPDF。"""
    settings = get_settings()
    if settings.use_baidu_pdf_parser:
        try:
            from app.services.baidu_doc_parser import parse_pdf as _baidu_parse_pdf

            text = _baidu_parse_pdf(str(path))
            if text and text.strip():
                logger.info("PDF 由百度·文档解析完成：%s（%d 字）", path.name, len(text))
                return text
            logger.warning("百度·文档解析返回空，回退本地解析：%s", path.name)
        except Exception as e:
            logger.warning("百度·文档解析失败，回退本地解析：%s（%s）", path.name, e)
    return _extract_pdf(path)


def _extract_pdf(path: Path) -> str:
    settings = get_settings()
    try:
        doc = fitz.open(str(path))
    except Exception as e:
        msg = str(e)
        code = "ERR_DOC_ENCRYPTED" if "password" in msg.lower() or "encrypt" in msg.lower() else "ERR_DOC_CORRUPTED"
        friendly = (
            "该文档处于加密保护，系统无法深度解析，请上传脱密版文档"
            if code == "ERR_DOC_ENCRYPTED"
            else f"PDF 无法打开，文件可能已损坏: {e}"
        )
        raise ParseError(friendly, code) from e
    try:
        if doc.page_count > settings.max_pdf_pages:
            raise ParseError(f"PDF 超过 {settings.max_pdf_pages} 页限制（Demo）")
        native_parts = _extract_pdf_native_text(doc)
        native_text = "\n\n".join(native_parts)
        native_cjk = cjk_ratio(native_text)
        native_low_quality = _is_low_quality_pdf_native_text(native_text)

        parts = native_parts
        if native_low_quality:
            # 原生抽取疑似乱码/低质量时，尝试 OCR 纠正。
            ocr_parts = _extract_pdf_ocr_text(doc)
            ocr_text = "\n\n".join(ocr_parts)
            ocr_cjk = cjk_ratio(ocr_text)
            # region agent log
            _debug_log(
                "pre-fix",
                "H36",
                "parser.py:_extract_pdf",
                "pdf native quality check and ocr fallback",
                {
                    "native_len": len(native_text),
                    "native_cjk_ratio": round(native_cjk, 4),
                    "native_low_quality": native_low_quality,
                    "ocr_len": len(ocr_text),
                    "ocr_cjk_ratio": round(ocr_cjk, 4),
                },
            )
            # endregion
            if ocr_parts and (ocr_cjk >= native_cjk + 0.02 or len(ocr_text) > len(native_text)):
                parts = ocr_parts
        elif not parts:
            # 扫描件兜底：尝试 OCR（需本机安装 tesseract + 中文语言包）。
            parts = _extract_pdf_ocr_text(doc)

        final_text = "\n\n".join(parts)
        final_text = _tidy_cjk_text(final_text)
        final_low_quality = _is_low_quality_pdf_native_text(final_text)
        # region agent log
        _debug_log(
            "pre-fix",
            "H37",
            "parser.py:_extract_pdf",
            "pdf final quality gate",
            {
                "final_len": len(final_text),
                "final_cjk_ratio": round(cjk_ratio(final_text), 4),
                "final_low_quality": final_low_quality,
                "chi_sim_available": _has_tesseract_lang("chi_sim"),
            },
        )
        # endregion

        if not parts or final_low_quality:
            if not _has_tesseract_lang("chi_sim"):
                raise ParseError(
                    "PDF 识别结果质量过低，且未检测到 Tesseract 中文语言包 chi_sim。"
                    " 请安装 chi_sim.traineddata 后在制度库点击“重新解析”。",
                    "ERR_DOC_LOW_QUALITY",
                )
            raise ParseError(
                "PDF 未能提取到有效文字（可能为扫描件）。"
                " 请安装 Tesseract OCR（含 chi_sim 中文包）后重试，"
                " 或先将 PDF 进行 OCR 后再上传。",
                "ERR_DOC_LOW_QUALITY",
            )
        # AI 排版改为按需触发（制度预览页里手动点），上传/解析只做确定性规整，保证上传快。
        return final_text
    finally:
        doc.close()


_CJK_RANGE = "一-鿿"
_CJK_CJK_SP = re.compile(rf"(?<=[{_CJK_RANGE}])[ \t]+(?=[{_CJK_RANGE}])")
_CJK_NUM_SP = re.compile(rf"(?<=[{_CJK_RANGE}])[ \t]+(?=[0-9])")
_NUM_CJK_SP = re.compile(rf"(?<=[0-9])[ \t]+(?=[{_CJK_RANGE}])")
_PUNCT_AFTER_SP = re.compile(rf"(?<=[{_CJK_RANGE}0-9])[ \t]+(?=[，。、；：？！）】」』])")
_PUNCT_BEFORE_SP = re.compile(rf"(?<=[（【「『])[ \t]+(?=[{_CJK_RANGE}0-9])")
_MULTI_SP = re.compile(r"[ \t]{2,}")

# 章节/条款级标题：作为“上一行”时不吸收下一行，保持独立成行。
_BLOCK_HEADING = re.compile(r"^(?:第[一二三四五六七八九十百千零〇\d]+[章节条款]|[一二三四五六七八九十]+、)")
# 新行起始标记：作为“当前行”时另起，不并入上一行。
_LINE_START = re.compile(
    r"^(?:"
    r"第[一二三四五六七八九十百千零〇\d]+[章节条款编]"
    r"|[（(][一二三四五六七八九十百\d]+[)）]"
    r"|[一二三四五六七八九十]+[、.]"
    r"|\d+[、.．]"
    r"|[A-Za-z][.)、）]"
    r"|[•·▪◆○※]"
    r")"
)
_SENT_END = "。！？；…!?;"


def _clean_inline_spaces(s: str) -> str:
    for _ in range(2):
        s = _CJK_CJK_SP.sub("", s)
    s = _CJK_NUM_SP.sub("", s)
    s = _NUM_CJK_SP.sub("", s)
    s = _PUNCT_AFTER_SP.sub("", s)
    s = _PUNCT_BEFORE_SP.sub("", s)
    s = _MULTI_SP.sub(" ", s)
    return s.strip()


def _tidy_cjk_text(text: str) -> str:
    """确定性规整中文 PDF 文本：删字间/标点边界多余空格、合并被硬断的句子、压缩空行。

    不依赖大模型，也不改动任何文字内容，只整理空格与换行。章节条款标题与列表
    编号行保持独立，避免破坏下游条款切分。
    """
    if not text:
        return text
    lines = [
        _clean_inline_spaces(ln)
        for ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    ]
    merged: list[str] = []
    for ln in lines:
        if not ln:
            if merged and merged[-1] != "":
                merged.append("")
            continue
        prev = merged[-1] if merged else ""
        can_merge = (
            bool(prev)
            and not _BLOCK_HEADING.match(prev)
            and prev[-1] not in _SENT_END
            and not _LINE_START.match(ln)
        )
        if can_merge:
            merged[-1] = prev + ln
        else:
            merged.append(ln)
    while merged and merged[-1] == "":
        merged.pop()
    return "\n".join(merged).strip()


def _extract_pdf_native_text(doc: fitz.Document) -> list[str]:
    parts: list[str] = []
    for page in doc:
        raw = page.get_text("text", sort=True).strip()
        cleaned = clean_page_body(raw)
        if cleaned:
            parts.append(cleaned)
    return parts


def _is_low_quality_pdf_native_text(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    # 有一定长度但几乎无中文，常见于编码错乱/提取异常。
    if len(t) >= 80 and cjk_ratio(t) < 0.03:
        return True
    # 明显的“字母/符号占绝对主导”且正文长度较长，判定为低质量。
    letters = sum(1 for ch in t if ("A" <= ch <= "Z") or ("a" <= ch <= "z"))
    non_space = sum(1 for ch in t if not ch.isspace())
    if non_space >= 200 and letters / max(non_space, 1) > 0.7 and cjk_ratio(t) < 0.05:
        return True
    return False


def _extract_pdf_ocr_text(doc: fitz.Document) -> list[str]:
    # 未安装 tesseract 时，直接跳过 OCR（保持错误提示可读）。
    tesseract = _resolve_tesseract_path()
    if not tesseract:
        return []
    tesseract_dir = str(Path(tesseract).parent)
    if tesseract_dir not in os.environ.get("PATH", ""):
        os.environ["PATH"] = tesseract_dir + os.pathsep + os.environ.get("PATH", "")

    # 优先使用项目内置语言包目录（可写，避免 Program Files 权限问题）。
    runtime_tessdata = _prepare_runtime_tessdata_prefix()
    if runtime_tessdata:
        # PyMuPDF/Tesseract 在当前环境下要求 TESSDATA_PREFIX 直接指向 tessdata 目录。
        os.environ["TESSDATA_PREFIX"] = str(runtime_tessdata)
        # region agent log
        _debug_log(
            "pre-fix",
            "H39",
            "parser.py:_extract_pdf_ocr_text",
            "set tessdata prefix",
            {
                    "runtime_prefix": os.environ.get("TESSDATA_PREFIX", ""),
                "runtime_tessdata_dir": str(runtime_tessdata.resolve()),
            },
        )
        # endregion

    parts: list[str] = []
    languages = ("chi_sim+eng", "chi_sim", "eng")
    page_logs: list[dict] = []
    for page in doc:
        text = ""
        used_lang = ""
        lang_try_logs: list[dict] = []
        for lang in languages:
            try:
                # PyMuPDF OCR：对扫描页生成 OCR TextPage 后再取文本。
                tp = page.get_textpage_ocr(language=lang, dpi=300, full=False)
                text = page.get_text("text", textpage=tp).strip()
                lang_try_logs.append(
                    {"lang": lang, "mode": "partial", "raw_len": len(text), "cjk_ratio": round(cjk_ratio(text), 4)}
                )
            except Exception as e:
                lang_try_logs.append({"lang": lang, "mode": "partial", "error": str(e)[:180]})
                text = ""
            if not text:
                try:
                    # 某些扫描页在 full=False 下取不到文字，回退 full=True 再试。
                    tp_full = page.get_textpage_ocr(language=lang, dpi=300, full=True)
                    text = page.get_text("text", textpage=tp_full).strip()
                    lang_try_logs.append(
                        {"lang": lang, "mode": "full", "raw_len": len(text), "cjk_ratio": round(cjk_ratio(text), 4)}
                    )
                except Exception as e:
                    lang_try_logs.append({"lang": lang, "mode": "full", "error": str(e)[:180]})
                    text = ""
            if text:
                used_lang = lang
                break
        cleaned = clean_page_body(text)
        if not cleaned and text and cjk_ratio(text) >= 0.02:
            # OCR 原文有一定中文但清洗后为空时，保底使用原文避免过清洗。
            cleaned = text
        if cleaned:
            parts.append(cleaned)
        page_logs.append(
            {
                "page": page.number + 1,
                "used_lang": used_lang,
                "raw_len": len(text or ""),
                "cleaned_len": len(cleaned or ""),
                "tries": lang_try_logs,
            }
        )
    # region agent log
    _debug_log(
        "pre-fix",
        "H38",
        "parser.py:_extract_pdf_ocr_text",
        "ocr extraction page diagnostics",
        {"pages": len(page_logs), "accepted_parts": len(parts), "detail": page_logs[:8]},
    )
    # endregion
    return parts


def _resolve_tesseract_path() -> str | None:
    candidates = [
        shutil.which("tesseract"),
        str(Path("C:/Program Files/Tesseract-OCR/tesseract.exe")),
        str(Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe")),
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def _has_tesseract_lang(lang: str) -> bool:
    candidates: list[Path] = []
    env_prefix = os.environ.get("TESSDATA_PREFIX", "").strip()
    if env_prefix:
        p = Path(env_prefix)
        candidates.append(p)
        candidates.append(p / "tessdata")
    candidates.extend(_candidate_tessdata_dirs())
    runtime_dir = _runtime_tessdata_dir()
    candidates.append(runtime_dir)
    candidates.extend(
        [
            Path("C:/Program Files/Tesseract-OCR/tessdata"),
            Path("C:/Program Files (x86)/Tesseract-OCR/tessdata"),
        ]
    )
    for base in candidates:
        if not base:
            continue
        f = base / f"{lang}.traineddata"
        if f.exists():
            return True
    return False


def _candidate_tessdata_dirs() -> list[Path]:
    return [
        ROOT_DIR / "data" / "tessdata",
        ROOT_DIR / "backend" / "data" / "tessdata",
    ]


def _runtime_tessdata_dir() -> Path:
    return Path(tempfile.gettempdir()) / "ic_tesseract_runtime" / "tessdata"


def _prepare_runtime_tessdata_prefix() -> Path | None:
    runtime_tessdata = _runtime_tessdata_dir()
    runtime_tessdata.mkdir(parents=True, exist_ok=True)
    copied = 0
    seen: set[str] = set()
    for src_dir in _candidate_tessdata_dirs() + [
        Path("C:/Program Files/Tesseract-OCR/tessdata"),
        Path("C:/Program Files (x86)/Tesseract-OCR/tessdata"),
    ]:
        if not src_dir.exists():
            continue
        for f in src_dir.glob("*.traineddata"):
            name = f.name.lower()
            if name in seen:
                continue
            dst = runtime_tessdata / f.name
            try:
                shutil.copy2(f, dst)
                seen.add(name)
                copied += 1
            except Exception:
                continue
    if copied == 0 and not any(runtime_tessdata.glob("*.traineddata")):
        return None
    return runtime_tessdata


def split_text_to_chunks(text: str, max_chars: int = 12000) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        start = end
    return chunks
